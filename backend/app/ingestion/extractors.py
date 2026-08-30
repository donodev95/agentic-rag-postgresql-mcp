"""In-memory PDF, DOCX, and UTF-8 text extractors."""

from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from backend.app.ingestion.types import ExtractedPage

SUPPORTED_MIME_TYPES = {
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    ".txt": {"text/plain"},
}


class DocumentValidationError(ValueError):
    """Raised when an upload is unsupported, empty, corrupt, or protected."""


def validate_upload(filename: str, mime_type: str, data: bytes, max_size_bytes: int) -> str:
    """Validate filename, declared media type, size, and basic file signature."""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_MIME_TYPES:
        raise DocumentValidationError("Supported file types are PDF, DOCX, and TXT")
    if mime_type.lower() not in SUPPORTED_MIME_TYPES[extension]:
        raise DocumentValidationError("File extension and MIME type do not match")
    if not data:
        raise DocumentValidationError("Uploaded file is empty")
    if len(data) > max_size_bytes:
        raise DocumentValidationError("Uploaded file exceeds the configured size limit")
    if extension == ".pdf" and not data.startswith(b"%PDF-"):
        raise DocumentValidationError("Invalid PDF signature")
    if extension == ".docx" and not data.startswith(b"PK"):
        raise DocumentValidationError("Invalid DOCX signature")
    if extension == ".txt" and b"\x00" in data:
        raise DocumentValidationError("TXT uploads cannot contain null bytes")
    return extension


def extract_pages(extension: str, data: bytes) -> list[ExtractedPage]:
    """Extract logical pages from validated in-memory bytes."""
    if extension == ".pdf":
        return _extract_pdf(data)
    if extension == ".docx":
        return _extract_docx(data)
    if extension == ".txt":
        return _extract_txt(data)
    raise DocumentValidationError("Unsupported file type")


def _extract_pdf(data: bytes) -> list[ExtractedPage]:
    try:
        reader = PdfReader(BytesIO(data), strict=True)
        if reader.is_encrypted:
            raise DocumentValidationError("Password-protected PDFs are not supported")
        pages = [
            ExtractedPage(text=page.extract_text() or "", page_number=index + 1)
            for index, page in enumerate(reader.pages)
        ]
    except DocumentValidationError:
        raise
    except (PdfReadError, OSError, ValueError) as exc:
        raise DocumentValidationError("PDF is corrupt or cannot be read") from exc
    if not any(page.text.strip() for page in pages):
        raise DocumentValidationError("PDF does not contain extractable text")
    return pages


def _extract_docx(data: bytes) -> list[ExtractedPage]:
    try:
        document = DocxDocument(BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    except (BadZipFile, KeyError, OSError, ValueError) as exc:
        raise DocumentValidationError("DOCX is corrupt or cannot be read") from exc
    text = "\n".join(part for part in parts if part.strip())
    if not text.strip():
        raise DocumentValidationError("DOCX does not contain extractable text")
    return [ExtractedPage(text=text, page_number=None)]


def _extract_txt(data: bytes) -> list[ExtractedPage]:
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentValidationError("TXT uploads must use UTF-8 encoding") from exc
    if not text.strip():
        raise DocumentValidationError("TXT does not contain text")
    return [ExtractedPage(text=text, page_number=None)]
